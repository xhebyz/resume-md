import sys
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def convert_html_to_word(html_content, output_file):
    doc = Document()
    soup = BeautifulSoup(html_content, 'html.parser')

    for tag in soup.find_all():
        if tag.name == 'h1':
            doc.add_heading(tag.text, level=1)
        elif tag.name == 'p':
            p = doc.add_paragraph(tag.text)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.save(output_file)
    print(f"Documento de Word '{output_file}' generado con éxito.")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Uso: python html_to_word.py archivo_html.html archivo_word.docx")
    else:
        input_html_file = sys.argv[1]
        output_word_file = sys.argv[2]

        with open(input_html_file, "r") as f:
            html_content = f.read()

        convert_html_to_word(html_content, output_word_file)
